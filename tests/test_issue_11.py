import io
import os
import shutil
import tempfile
import unittest
from pathlib import Path

TEST_ROOT = Path(__file__).resolve().parents[1] / ".tmp_validation" / "issue11_tests"
TEST_ROOT.mkdir(parents=True, exist_ok=True)

os.environ.setdefault("DATABASE_PATH", str(TEST_ROOT / "bootstrap.db"))
os.environ.setdefault("DOWNLOAD_DIR", str(TEST_ROOT / "bootstrap_uploads"))

import app as app_module
from docx import Document
from document_processing import (
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    build_chunk_records,
    extract_plain_text,
    process_uploaded_file,
    split_text,
)


def write_text_pdf(filepath: Path, text: str) -> None:
    objects = []

    def add_object(body: str) -> int:
        objects.append(body)
        return len(objects)

    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    catalog_id = add_object("<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_object("<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    page_id = add_object(
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] "
        "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
    )
    content_stream = f"BT\n/F1 18 Tf\n36 100 Td\n({escaped_text}) Tj\nET"
    content_id = add_object(
        f"<< /Length {len(content_stream.encode('latin-1'))} >>\nstream\n{content_stream}\nendstream"
    )
    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    assert [catalog_id, pages_id, page_id, content_id, font_id] == [1, 2, 3, 4, 5]

    pdf_parts = ["%PDF-1.4\n"]
    offsets = [0]

    for index, body in enumerate(objects, start=1):
        offsets.append(sum(len(part.encode("latin-1")) for part in pdf_parts))
        pdf_parts.append(f"{index} 0 obj\n{body}\nendobj\n")

    xref_offset = sum(len(part.encode("latin-1")) for part in pdf_parts)
    pdf_parts.append(f"xref\n0 {len(objects) + 1}\n")
    pdf_parts.append("0000000000 65535 f \n")

    for offset in offsets[1:]:
        pdf_parts.append(f"{offset:010d} 00000 n \n")

    pdf_parts.append(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    )

    filepath.write_bytes("".join(pdf_parts).encode("latin-1"))


class Issue11ImplementationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.base_dir = TEST_ROOT / "runtime"
        cls.base_dir.mkdir(parents=True, exist_ok=True)
        cls.database_path = cls.base_dir / "test.db"
        cls.download_dir = cls.base_dir / "uploads"
        cls.download_dir.mkdir(parents=True, exist_ok=True)

        app_module.DATABASE = str(cls.database_path)
        app_module.DOWNLOAD_DIR = str(cls.download_dir)
        os.makedirs(app_module.DOWNLOAD_DIR, exist_ok=True)
        app_module.app.config["TESTING"] = True

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.base_dir, ignore_errors=True)

    def setUp(self):
        if self.database_path.exists():
            self.database_path.unlink()

        if self.download_dir.exists():
            shutil.rmtree(self.download_dir)

        self.download_dir.mkdir(parents=True, exist_ok=True)
        app_module.CHUNK_REGISTRY.clear()
        app_module.UPLOAD_CHUNK_INDEX.clear()
        app_module.init_db()
        self.client = app_module.app.test_client()

    def login_as_admin(self):
        with self.client.session_transaction() as session:
            session["user"] = "admin"
            session["role"] = "admin"
            session["selected_course"] = "Ingenieria de Software"
            session["session_id"] = "test-session"

    def create_temp_file(self, suffix: str, content: str) -> Path:
        path = self.base_dir / f"sample{suffix}"
        path.write_text(content, encoding="utf-8")
        return path

    def test_extract_plain_text_for_text_markdown_docx_and_pdf(self):
        txt_path = self.create_temp_file(".txt", "texto plano")
        md_path = self.create_temp_file(".md", "# titulo\n\ncontenido markdown")

        docx_path = self.base_dir / "sample.docx"
        document = Document()
        document.add_paragraph("Primera linea")
        document.add_paragraph("Segunda linea")
        document.save(docx_path)

        pdf_path = self.base_dir / "sample.pdf"
        write_text_pdf(pdf_path, "Texto PDF")

        self.assertEqual(extract_plain_text(str(txt_path)), "texto plano")
        self.assertIn("contenido markdown", extract_plain_text(str(md_path)))
        self.assertEqual(extract_plain_text(str(docx_path)), "Primera linea\nSegunda linea")
        self.assertIn("Texto PDF", extract_plain_text(str(pdf_path)))

    def test_split_text_uses_expected_chunk_size_and_overlap(self):
        text = "A" * (DEFAULT_CHUNK_SIZE + 200)
        chunks = split_text(text)

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(len(chunk) <= DEFAULT_CHUNK_SIZE for chunk in chunks))
        self.assertEqual(chunks[0][-DEFAULT_CHUNK_OVERLAP:], chunks[1][:DEFAULT_CHUNK_OVERLAP])

    def test_build_chunk_records_uses_deterministic_ids(self):
        records = build_chunk_records(
            ["chunk one", "chunk two"],
            document_id=7,
            doc_hash="doc123",
            upload_hash="up456",
            course="Ingenieria de Software",
            upload_date="2026-03-22 10:00:00",
            filename="sample.txt",
            file_hash="file789",
        )

        self.assertEqual(records[0]["chunk_id"], "up456:7:0")
        self.assertEqual(records[1]["chunk_id"], "up456:7:1")
        self.assertEqual(records[1]["doc_hash"], "doc123")

    def test_process_uploaded_file_returns_empty_list_for_blank_text(self):
        blank_path = self.create_temp_file(".txt", "   \n")
        records = process_uploaded_file(
            str(blank_path),
            document_id=1,
            doc_hash="doc1",
            upload_hash="upload1",
            course="Ingenieria de Software",
            upload_date="2026-03-22 10:00:00",
            filename="blank.txt",
            file_hash="file1",
        )

        self.assertEqual(records, [])

    def test_upload_generates_chunks_and_keeps_registry_in_memory(self):
        self.login_as_admin()

        response = self.client.post(
            "/api/upload",
            data={"files[]": (io.BytesIO(b"contenido " * 200), "notas.txt")},
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        upload_hash = payload["upload_hash"]
        document_id = payload["files"][0]["document_id"]

        self.assertIn(upload_hash, app_module.UPLOAD_CHUNK_INDEX)
        chunk_ids = app_module.UPLOAD_CHUNK_INDEX[upload_hash]
        self.assertGreater(len(chunk_ids), 0)
        first_chunk = app_module.CHUNK_REGISTRY[chunk_ids[0]]
        self.assertEqual(first_chunk["document_id"], document_id)
        self.assertEqual(first_chunk["upload_hash"], upload_hash)
        self.assertEqual(first_chunk["filename"], "notas.txt")

    def test_upload_reuses_doc_hash_for_new_versions(self):
        self.login_as_admin()

        first_response = self.client.post(
            "/api/upload",
            data={"files[]": (io.BytesIO(b"version uno " * 120), "versionado.txt")},
            content_type="multipart/form-data",
        )
        second_response = self.client.post(
            "/api/upload",
            data={"files[]": (io.BytesIO(b"version dos " * 120), "versionado.txt")},
            content_type="multipart/form-data",
        )

        first_payload = first_response.get_json()
        second_payload = second_response.get_json()

        self.assertEqual(first_response.status_code, 200)
        self.assertEqual(second_response.status_code, 200)
        self.assertEqual(first_payload["files"][0]["document_id"], second_payload["files"][0]["document_id"])
        self.assertEqual(first_payload["files"][0]["doc_hash"], second_payload["files"][0]["doc_hash"])

    def test_multi_file_upload_shares_one_upload_hash(self):
        self.login_as_admin()

        response = self.client.post(
            "/api/upload",
            data={
                "files[]": [
                    (io.BytesIO(b"alpha " * 200), "alpha.txt"),
                    (io.BytesIO(b"beta " * 200), "beta.md"),
                ]
            },
            content_type="multipart/form-data",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        upload_hash = payload["upload_hash"]
        chunk_ids = app_module.UPLOAD_CHUNK_INDEX[upload_hash]

        self.assertEqual(len(payload["files"]), 2)
        self.assertTrue(chunk_ids)
        self.assertTrue(all(chunk_id.startswith(f"{upload_hash}:") for chunk_id in chunk_ids))

    def test_documents_endpoint_still_returns_uploaded_document(self):
        self.login_as_admin()

        upload_response = self.client.post(
            "/api/upload",
            data={"files[]": (io.BytesIO(b"regresion " * 50), "regresion.txt")},
            content_type="multipart/form-data",
        )
        self.assertEqual(upload_response.status_code, 200)

        documents_response = self.client.get("/api/documents/Ingenieria%20de%20Software")
        self.assertEqual(documents_response.status_code, 200)
        documents = documents_response.get_json()

        self.assertEqual(len(documents), 1)
        self.assertEqual(documents[0]["filename"], "regresion.txt")


if __name__ == "__main__":
    unittest.main()
